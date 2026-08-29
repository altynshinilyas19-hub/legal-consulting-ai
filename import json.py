import json
import re
import time
from pathlib import Path
from urllib.parse import urljoin, urlparse

from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


START_URL = "https://demo.consultant.ru/cgi/online.cgi?req=query&mode=fullsplus&histid=9434268&cacheid=E11929F7CF548DC0BE91D1D5E7A55A1B&ts=CD9SCEVi1JMkR8zj&rnd=BTrJVA#TH9SCEVCw4mE8ZBX1"

OUT_DIR = Path("consultant_output")
STATE_FILE = OUT_DIR / "state.json"
TEXTS_DIR = OUT_DIR / "cases_txt"
DOCX_DIR = OUT_DIR / "cases_docx"
JSONL_FILE = OUT_DIR / "all_cases.jsonl"
SEEN_FILE = OUT_DIR / "seen_urls.txt"

OUT_DIR.mkdir(exist_ok=True)
TEXTS_DIR.mkdir(exist_ok=True)
DOCX_DIR.mkdir(exist_ok=True)


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()


def safe_filename(text: str, max_len: int = 120) -> str:
    text = norm(text)
    text = re.sub(r'[\\/*?:"<>|]', "_", text)
    text = text.strip(" .")
    if not text:
        text = "case"
    return text[:max_len]


def load_seen_urls() -> set:
    if not SEEN_FILE.exists():
        return set()
    return set(line.strip() for line in SEEN_FILE.read_text(encoding="utf-8").splitlines() if line.strip())


def append_seen_url(url: str):
    with open(SEEN_FILE, "a", encoding="utf-8") as f:
        f.write(url + "\n")


def append_jsonl(row: dict):
    with open(JSONL_FILE, "a", encoding="utf-8") as f:
        f.write(json.dumps(row, ensure_ascii=False) + "\n")


def save_txt_case(row: dict, index_num: int):
    title = row.get("title") or f"case_{index_num}"
    fname = f"{index_num:06d}_{safe_filename(title)}.txt"
    path = TEXTS_DIR / fname

    content = []
    content.append(f"TITLE: {row.get('title', '')}")
    content.append(f"URL: {row.get('url', '')}")
    content.append(f"PAGE: {row.get('page_num', '')}")
    content.append("")
    content.append(row.get("text_raw", ""))

    path.write_text("\n".join(content), encoding="utf-8")
    return path


def save_docx_case(row: dict, index_num: int):
    if not DOCX_AVAILABLE:
        return None

    title = row.get("title") or f"case_{index_num}"
    fname = f"{index_num:06d}_{safe_filename(title)}.docx"
    path = DOCX_DIR / fname

    doc = Document()
    doc.add_heading(row.get("title", "Без названия"), level=1)
    doc.add_paragraph(f"URL: {row.get('url', '')}")
    doc.add_paragraph(f"Страница выдачи: {row.get('page_num', '')}")
    doc.add_paragraph("")
    doc.add_paragraph(row.get("text_raw", ""))

    doc.save(str(path))
    return path


def extract_text_by_selectors(page, selectors):
    for sel in selectors:
        try:
            loc = page.locator(sel).first
            if loc.count() > 0:
                txt = norm(loc.inner_text(timeout=3000))
                if txt:
                    return txt
        except Exception:
            pass
    return ""


def auto_scroll(page, steps=8, pause=1.0):
    for _ in range(steps):
        try:
            page.mouse.wheel(0, 2000)
        except Exception:
            pass
        try:
            page.evaluate("window.scrollBy(0, document.body.scrollHeight)")
        except Exception:
            pass
        time.sleep(pause)


def collect_result_links(page):
    links = []
    seen = set()

    selectors = [
        "a[href*='online.cgi']",
        "a[href*='req=doc']",
        "a[href*='req=home']",
        "a[href*='req=query']",
        ".result a",
        ".search-results a",
        ".docs a",
        "a",
    ]

    for sel in selectors:
        try:
            nodes = page.locator(sel)
            count = min(nodes.count(), 1500)

            for i in range(count):
                try:
                    node = nodes.nth(i)
                    href = node.get_attribute("href")
                    text = norm(node.inner_text(timeout=700))

                    if not href:
                        continue

                    full_url = urljoin(page.url, href)

                    if "demo.consultant.ru" not in full_url:
                        continue
                    if "online.cgi" not in full_url:
                        continue
                    if full_url in seen:
                        continue

                    # выкидываем слишком мусорные короткие пункты
                    if len(text) < 5:
                        continue

                    seen.add(full_url)
                    links.append({
                        "title_hint": text,
                        "url": full_url,
                    })
                except Exception:
                    pass

            if links:
                break
        except Exception:
            pass

    return links


def extract_case_data(page, url: str):
    try:
        page.wait_for_load_state("domcontentloaded", timeout=10000)
    except Exception:
        pass

    try:
        page.wait_for_load_state("networkidle", timeout=8000)
    except Exception:
        pass

    time.sleep(1.5)

    title = extract_text_by_selectors(page, [
        "h1",
        ".document-title",
        ".docHeader",
        ".title",
        "title",
        "body",
    ])[:500]

    text_raw = extract_text_by_selectors(page, [
        ".document",
        ".docText",
        ".content",
        "#content",
        ".texto",
        "body",
    ])

    return {
        "url": url,
        "title": title,
        "text_raw": text_raw,
    }


def looks_like_real_case(row: dict, min_text_len: int = 500) -> bool:
    text = row.get("text_raw", "")
    title = (row.get("title") or "").lower()

    if len(text) < min_text_len:
        return False

    bad_words = [
        "консультантплюс",
        "поиск",
        "расширенный поиск",
        "следующая",
        "предыдущая",
    ]

    if any(word in title for word in bad_words) and len(text) < 2000:
        return False

    return True


def click_next_page(page) -> bool:
    candidates = [
        "a:has-text('Следующая')",
        "button:has-text('Следующая')",
        "a:has-text('Далее')",
        "button:has-text('Далее')",
        "a[rel='next']",
        ".pager a",
        ".pagination a",
        "a",
        "button",
    ]

    for sel in candidates:
        try:
            nodes = page.locator(sel)
            count = min(nodes.count(), 200)

            for i in range(count):
                try:
                    node = nodes.nth(i)
                    text = norm(node.inner_text(timeout=500)).lower()
                    href = node.get_attribute("href")

                    good_by_text = text in {"следующая", "далее", ">", ">>"} or "следующ" in text or "далее" in text
                    good_by_href = href and ("page=" in href or "start=" in href or "next" in href.lower())

                    if not good_by_text and not good_by_href:
                        continue

                    node.scroll_into_view_if_needed(timeout=2000)
                    time.sleep(0.5)
                    node.click(timeout=4000)

                    try:
                        page.wait_for_load_state("domcontentloaded", timeout=10000)
                    except Exception:
                        pass
                    try:
                        page.wait_for_load_state("networkidle", timeout=8000)
                    except Exception:
                        pass

                    time.sleep(1.5)
                    return True
                except Exception:
                    pass
        except Exception:
            pass

    return False


def bootstrap_login():
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=False)
        context = browser.new_context()
        page = context.new_page()

        page.goto(START_URL, wait_until="domcontentloaded")

        print("\n[1] Войди в аккаунт вручную")
        print("[2] Открой страницу, где виден список дел")
        print("[3] Когда список дел уже открыт — нажми Enter в консоли\n")
        input()

        context.storage_state(path=str(STATE_FILE))
        browser.close()
        print(f"[OK] Сессия сохранена: {STATE_FILE}")


def scrape_cases(max_pages=50, min_text_len=500):
    if not STATE_FILE.exists():
        raise RuntimeError("Сначала запусти bootstrap_login() и сохрани сессию.")

    seen_urls = load_seen_urls()
    saved_count = len(seen_urls)

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=False, slow_mo=100)
        context = browser.new_context(storage_state=str(STATE_FILE))
        page = context.new_page()

        page.goto(START_URL, wait_until="domcontentloaded")
        try:
            page.wait_for_load_state("networkidle", timeout=8000)
        except Exception:
            pass

        for page_num in range(1, max_pages + 1):
            print(f"\n========== PAGE {page_num} ==========")

            auto_scroll(page, steps=10, pause=1.0)
            links = collect_result_links(page)

            print(f"Найдено ссылок на странице: {len(links)}")

            page_saved = 0

            for item in links:
                url = item["url"]

                if url in seen_urls:
                    continue

                detail = context.new_page()
                try:
                    print(f"OPEN: {url}")
                    detail.goto(url, wait_until="domcontentloaded", timeout=20000)

                    row = extract_case_data(detail, url)
                    row["page_num"] = page_num

                    print(f"TITLE: {row['title'][:150]}")
                    print(f"TEXT LEN: {len(row['text_raw'])}")

                    if not looks_like_real_case(row, min_text_len=min_text_len):
                        print("SKIP: мало текста или это не карточка дела")
                        continue

                    saved_count += 1
                    page_saved += 1

                    append_jsonl(row)
                    save_txt_case(row, saved_count)
                    save_docx_case(row, saved_count)
                    append_seen_url(url)
                    seen_urls.add(url)

                    print(f"SAVED: #{saved_count}")

                except PlaywrightTimeoutError:
                    print("TIMEOUT")
                except Exception as e:
                    print(f"ERROR: {e}")
                finally:
                    detail.close()

                    time.sleep(0.5)

            print(f"Сохранено на этой странице: {page_saved}")

            moved = click_next_page(page)
            if not moved:
                print("Следующая страница не найдена. Останавливаюсь.")
                break

        browser.close()

    print(f"\nГотово. Всего сохранено: {saved_count}")
    print(f"TXT файлы: {TEXTS_DIR}")
    if DOCX_AVAILABLE:
        print(f"DOCX файлы: {DOCX_DIR}")
    print(f"JSONL файл: {JSONL_FILE}")


if __name__ == "__main__":
    
    bootstrap_login()

    #scrape_cases(max_pages=100, min_text_len=400)


    
    
    
    
    
    
    
    
    
    
    
if __name__ == "__main__":
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=False)
        context = browser.new_context(storage_state="consultant_output/state.json")
        page = context.new_page()
        page.goto("https://demo.consultant.ru")
        input("Браузер открыт, нажми Enter чтобы закрыть...")